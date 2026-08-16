from __future__ import annotations

import gzip
import hashlib
import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from ai_education.services.oss_quick_diagnostic_bank import (
    StructuredOssQuickDiagnosticBank,
)
from ai_education.services.quick_diagnostic_bank import QuickDiagnosticBank
from scripts.build_oss_quick_diagnostic_bank import (
    build_question,
    docx_has_answer_marker,
    extract_source_questions,
)

PREFIX = "knowledge/processed/quick_diagnostic/v1"


def structured_question(index: int, *, subject: str = "mathematics") -> dict:
    return {
        "knowledge_focus": "集合与常用逻辑用语",
        "difficulty": 0.45,
        "prompt": f"集合与充分条件诊断题 {index}",
        "prompt_html": f"集合与充分条件诊断题 {index}",
        "options": [f"{letter} 选项 {index}" for letter in "ABCD"],
        "options_html": [f"{letter} 选项 {index}" for letter in "ABCD"],
        "correct_option": index % 4,
        "explanation": "依据集合关系与充分条件的定义判断。",
        "source_question_id": f"oss_{subject}_{index}",
        "source_paper_id": f"oss_paper_{subject}_{index}",
        "source_title": "授权教材配套练习",
        "source_document_sha256": "a" * 64,
        "source_storage": "oss",
        "source_kind": "licensed_textbook_practice",
        "knowledge_tags": ["集合", "充分条件"],
        "search_text": "集合 常用逻辑用语 充分条件 必要条件",
    }


def object_payloads(questions: list[dict]) -> dict[str, bytes]:
    shard_key = f"{PREFIX}/subjects/mathematics.json.gz"
    shard = gzip.compress(
        json.dumps(
            {
                "schema_version": "1.0",
                "subject": "mathematics",
                "questions": questions,
            },
            ensure_ascii=False,
        ).encode("utf-8"),
        mtime=0,
    )
    manifest = json.dumps(
        {
            "schema_version": "1.0",
            "subjects": {
                "mathematics": {
                    "object_key": shard_key,
                    "sha256": hashlib.sha256(shard).hexdigest(),
                    "question_count": len(questions),
                }
            },
        }
    ).encode()
    return {f"{PREFIX}/manifest.json": manifest, shard_key: shard}


class FakeReader:
    def __init__(self, objects: dict[str, bytes], *, fail: bool = False) -> None:
        self.objects = objects
        self.fail = fail
        self.calls: list[str] = []

    def get_bytes(self, key: str, *, max_bytes: int) -> bytes:
        self.calls.append(key)
        if self.fail:
            raise RuntimeError("temporary OSS outage")
        payload = self.objects[key]
        if len(payload) > max_bytes:
            raise ValueError("too large")
        return payload


class FakeSupplementalBank:
    def __init__(self, questions: list[dict]) -> None:
        self._questions = questions

    def questions(self, subject: str) -> list[dict]:
        return list(self._questions) if subject == "mathematics" else []


class StructuredOssQuestionBankTests(unittest.TestCase):
    def create_bank(self, cache_dir: Path, reader: FakeReader) -> StructuredOssQuickDiagnosticBank:
        return StructuredOssQuickDiagnosticBank(
            bucket="test-private-bucket",
            region="cn-hangzhou",
            endpoint="https://oss-cn-hangzhou-internal.aliyuncs.com",
            ecs_role_name="test-read-role",
            prefix=PREFIX,
            cache_dir=cache_dir,
            reader=reader,
        )

    def test_reads_checksum_verified_gzip_shard(self) -> None:
        questions = [structured_question(index) for index in range(10)]
        with tempfile.TemporaryDirectory() as temporary:
            reader = FakeReader(object_payloads(questions))
            loaded = self.create_bank(Path(temporary), reader).questions("mathematics")

        self.assertEqual(len(loaded), 10)
        self.assertTrue(all(item["source_storage"] == "oss" for item in loaded))
        self.assertEqual(
            reader.calls,
            [
                f"{PREFIX}/manifest.json",
                f"{PREFIX}/subjects/mathematics.json.gz",
            ],
        )

    def test_transient_outage_uses_last_verified_disk_cache(self) -> None:
        questions = [structured_question(index) for index in range(10)]
        with tempfile.TemporaryDirectory() as temporary:
            cache_dir = Path(temporary)
            self.create_bank(cache_dir, FakeReader(object_payloads(questions))).questions(
                "mathematics"
            )
            recovered = self.create_bank(cache_dir, FakeReader({}, fail=True)).questions(
                "mathematics"
            )

        self.assertEqual(len(recovered), 10)

    def test_rejects_tampered_shard(self) -> None:
        questions = [structured_question(index) for index in range(10)]
        objects = object_payloads(questions)
        manifest_key = f"{PREFIX}/manifest.json"
        manifest = json.loads(objects[manifest_key])
        manifest["subjects"]["mathematics"]["sha256"] = "0" * 64
        objects[manifest_key] = json.dumps(manifest).encode()
        with (
            tempfile.TemporaryDirectory() as temporary,
            self.assertRaisesRegex(ValueError, "校验和"),
        ):
            self.create_bank(Path(temporary), FakeReader(objects)).questions("mathematics")

    def test_rejects_gzip_shard_over_uncompressed_limit(self) -> None:
        questions = [structured_question(index) for index in range(10)]
        with (
            tempfile.TemporaryDirectory() as temporary,
            patch(
                "ai_education.services.oss_quick_diagnostic_bank.MAX_UNCOMPRESSED_SHARD_BYTES",
                128,
            ),
            self.assertRaisesRegex(ValueError, "解压后超过允许大小"),
        ):
            self.create_bank(Path(temporary), FakeReader(object_payloads(questions))).questions(
                "mathematics"
            )

    def test_quick_diagnostic_selects_oss_questions_by_chapter(self) -> None:
        questions = [structured_question(index) for index in range(12)]
        with tempfile.TemporaryDirectory() as temporary:
            bank = QuickDiagnosticBank(
                Path(temporary), supplemental_bank=FakeSupplementalBank(questions)
            )
            selected = bank.questions(
                subject="mathematics",
                seed="oss-chapter-selection",
                progress_label="必修 第一册 集合与常用逻辑用语",
                whole_book=False,
                scope_units=[
                    {
                        "id": "TB-MATH-SETS-C01",
                        "label": "必修 第一册 集合与常用逻辑用语",
                    }
                ],
            )

        self.assertEqual(len(selected), 10)
        self.assertTrue(
            all(
                item["provenance"]["source_storage"] == "oss"
                and item["provenance"]["source_kind"] == "licensed_textbook_practice"
                and item["provenance"]["scope_match_verified"]
                for item in selected
            )
        )

    def test_offline_builder_accepts_only_source_answered_choice(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            source_path = Path(temporary) / "集合与常用逻辑用语练习（带答案）.docx"
            document = Document()
            document.add_paragraph(
                "1. 已知集合 A={1,2}，下列说法正确的是 A. 1∈A B. 3∈A C. A 为空集 D. A 只有一个元素"
            )
            document.add_paragraph("【答案】 A")
            document.add_paragraph("【解析】 根据集合元素与集合的关系可知 1∈A。")
            document.save(source_path)

            payload = source_path.read_bytes()
            extracted = extract_source_questions("mathematics", source_path)
            built = build_question(
                extracted[0],
                subject="mathematics",
                relative_path="A版/数学/集合与常用逻辑用语练习（带答案）.docx",
            )

        self.assertTrue(docx_has_answer_marker(payload))
        self.assertEqual(len(extracted), 1)
        self.assertIsNotNone(built)
        assert built is not None
        self.assertEqual(built["correct_option"], 0)
        self.assertEqual(len(built["options"]), 4)
        self.assertNotIn("source_path", built)


if __name__ == "__main__":
    unittest.main()
